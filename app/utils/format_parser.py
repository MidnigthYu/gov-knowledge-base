"""多格式文档解析模块

提供 PDF、DOCX、TXT、Markdown 四类文档的文本提取能力，
对外暴露统一入口 extract_text，解析失败统一返回 None。
"""
import os
from typing import Optional
import pdfplumber
from docx import Document
from .text_cleaner import safe_read_file


def extract_text(file_path: str) -> Optional[str]:
    """统一格式解析入口，所有格式最终输出纯文本，解析失败返回 None"""
    ext = os.path.splitext(file_path)[1].lower()

    if ext in (".txt", ".md"):
        return safe_read_file(file_path)
    elif ext == ".pdf":
        return _extract_pdf_text(file_path)
    elif ext == ".docx":
        return _extract_docx_text(file_path)
    else:
        return None


def _extract_pdf_text(file_path: str) -> Optional[str]:
    """提取 PDF 全文文本，逐页拼接，解析失败返回 None"""
    try:
        text_list = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_list.append(page_text)
        return "\n".join(text_list)
    except Exception:
        return None


def _extract_docx_text(file_path: str) -> Optional[str]:
    """提取 DOCX 正文段落与表格文本，解析失败返回 None"""
    try:
        doc = Document(file_path)
        text_list = []

        # 提取正文段落
        for para in doc.paragraphs:
            if para.text.strip():
                text_list.append(para.text)

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_list.append(row_text)

        return "\n".join(text_list)
    except Exception:
        return None
